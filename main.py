# from fastapi import FastAPI, UploadFile, File, HTTPException, Query
# from fastapi.responses import JSONResponse
# from sentence_transformers import SentenceTransformer, util
# import faiss
# import numpy as np
# import os
# import shutil
# import uuid
# import PyPDF2
# import docx
# import json

# app = FastAPI()

# # Load Embedding Model
# try:
#     embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
# except Exception as e:
#     raise RuntimeError(f"Failed to load embedding model: {str(e)}")

# # Constants
# UPLOAD_DIR = "uploaded_files"
# INDEX_FILE = "faiss_index"
# DATA_FILE = "data.json"

# # Ensure Upload Directory Exists
# os.makedirs(UPLOAD_DIR, exist_ok=True)

# # FAISS Index Setup
# d = 384  # Embedding dimension
# index = faiss.IndexFlatL2(d) if not os.path.exists(INDEX_FILE) else faiss.read_index(INDEX_FILE)

# # Load Metadata
# if os.path.exists(DATA_FILE):
#     with open(DATA_FILE, "r") as f:
#         metadata_store = json.load(f)
# else:
#     metadata_store = {}

# def extract_text(file_path: str) -> str:
#     """Extracts text from PDF or DOCX files."""
#     _, ext = os.path.splitext(file_path)
    
#     if ext.lower() == ".pdf":
#         try:
#             with open(file_path, "rb") as f:
#                 reader = PyPDF2.PdfReader(f)
#                 return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
#         except Exception as e:
#             raise HTTPException(status_code=500, detail=f"Error reading PDF: {str(e)}")
    
#     elif ext.lower() == ".docx":
#         try:
#             doc = docx.Document(file_path)
#             return " ".join([para.text for para in doc.paragraphs])
#         except Exception as e:
#             raise HTTPException(status_code=500, detail=f"Error reading DOCX: {str(e)}")
    
#     raise HTTPException(status_code=400, detail="Unsupported file format")

# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     """Handles file uploads and stores embeddings in FAISS index."""
#     file_ext = os.path.splitext(file.filename)[1].lower()
#     if file_ext not in [".pdf", ".docx"]:
#         raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

#     file_id = str(uuid.uuid4())
#     file_path = os.path.join(UPLOAD_DIR, f"{file_id}{file_ext}")

#     try:
#         with open(file_path, "wb") as f:
#             shutil.copyfileobj(file.file, f)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")

#     text = extract_text(file_path)
#     if not text.strip():
#         os.remove(file_path)  
#         raise HTTPException(status_code=400, detail="Document contains no readable text")

#     try:
#         embedding = embedding_model.encode([text])[0]
#         index.add(np.array([embedding], dtype=np.float32))
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to generate embedding: {str(e)}")

#     metadata_store[file_id] = {"filename": file.filename, "content": text}

#     faiss.write_index(index, INDEX_FILE)
#     with open(DATA_FILE, "w") as f:
#         json.dump(metadata_store, f)

#     return JSONResponse(content={"message": "File uploaded successfully", "file_id": file_id})

# @app.get("/search/")
# async def search_documents(query: str = Query(..., min_length=3), top_k: int = Query(5, ge=1, le=20)):
#     """Searches for relevant documents based on user query."""
#     if index.ntotal == 0:
#         raise HTTPException(status_code=400, detail="No documents available for search")

#     try:
#         query_embedding = embedding_model.encode([query])[0].reshape(1, -1)
#         distances, indices = index.search(query_embedding, top_k)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

#     results = [
#         {"file_id": list(metadata_store.keys())[idx], "filename": metadata_store[list(metadata_store.keys())[idx]]["filename"], "distance": float(dist)}
#         for dist, idx in zip(distances[0], indices[0]) if idx < len(metadata_store)
#     ]

#     return JSONResponse(content={"results": results})

# @app.get("/index_stats/")
# async def index_stats():
#     """Returns the number of documents stored in FAISS index."""
#     return JSONResponse(content={"total_documents": index.ntotal})

# @app.post("/rebuild_index/")
# async def rebuild_index():
#     """Rebuilds the FAISS index from stored metadata and documents."""
#     global index
#     index = faiss.IndexFlatL2(d) 

#     if not metadata_store:
#         return JSONResponse(content={"message": "No documents found to rebuild index"})

#     try:
#         for file_id, data in metadata_store.items():
#             embedding = embedding_model.encode([data["content"]])[0]
#             index.add(np.array([embedding], dtype=np.float32))

#         faiss.write_index(index, INDEX_FILE)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to rebuild index: {str(e)}")

#     return JSONResponse(content={"message": "Index rebuilt successfully", "total_documents": index.ntotal})

# @app.delete("/delete/{file_id}")
# async def delete_document(file_id: str):
#     """Deletes a document from storage and FAISS index."""
#     if file_id not in metadata_store:
#         raise HTTPException(status_code=404, detail="File not found")

#     del metadata_store[file_id]

#     try:
#         with open(DATA_FILE, "w") as f:
#             json.dump(metadata_store, f)
#         faiss.write_index(index, INDEX_FILE)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")

#     return JSONResponse(content={"message": "File deleted successfully"})

# @app.post("/reset_index/")
# async def reset_index():
#     """Resets the FAISS index, deletes metadata, and clears uploaded files."""
#     global index, metadata_store

#     try:
#         index = faiss.IndexFlatL2(d)
#         faiss.write_index(index, INDEX_FILE)
#         metadata_store.clear()

#         if os.path.exists(DATA_FILE):
#             os.remove(DATA_FILE)

#         for file in os.listdir(UPLOAD_DIR):
#             file_path = os.path.join(UPLOAD_DIR, file)
#             if os.path.isfile(file_path):
#                 os.remove(file_path)

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Failed to reset index: {str(e)}")

#     return JSONResponse(content={"message": "Index reset successfully, all files and metadata deleted"})

# @app.get("/query/")
# async def query_documents(query: str = Query(..., min_length=3)):
#     """Fetches the **most relevant part** of a document instead of the full content."""
#     if not metadata_store:
#         raise HTTPException(status_code=404, detail="No documents available for search")

#     query_embedding = embedding_model.encode([query])[0].reshape(1, -1)
#     distances, indices = index.search(query_embedding, 5)

#     best_match, best_score, best_file = None, -1, None

#     for dist, idx in zip(distances[0], indices[0]):
#         if idx < len(metadata_store):
#             file_id = list(metadata_store.keys())[idx]
#             content = metadata_store[file_id]["content"]
#             sentences = content.split(". ")
#             sentence_embeddings = embedding_model.encode(sentences)
            
#             similarities = util.pytorch_cos_sim(query_embedding, sentence_embeddings)[0].cpu().numpy()
#             best_sentence = sentences[np.argmax(similarities)]

#             if similarities[np.argmax(similarities)] > best_score:
#                 best_score, best_match, best_file = similarities[np.argmax(similarities)], best_sentence, metadata_store[file_id]["filename"]

#     return JSONResponse(content={"file": best_file, "matched_text": best_match})


from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.responses import JSONResponse
from sentence_transformers import SentenceTransformer, util
import faiss
import numpy as np
import os
import shutil
import uuid
import PyPDF2
import docx
import json
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
import torch

app = FastAPI()

# Load Embedding Model
try:
    embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
except Exception as e:
    raise RuntimeError(f"Failed to load embedding model: {str(e)}")

# Load Phi-2 (2.7B) Model
try:
    model_name = "microsoft/phi-2"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForCausalLM.from_pretrained(model_name, torch_dtype=torch.float16)
    qa_pipeline = pipeline("text-generation", model=model, tokenizer=tokenizer)
except Exception as e:
    raise RuntimeError(f"Failed to load Phi-2 model: {str(e)}")

# Constants
UPLOAD_DIR = "uploaded_files"
INDEX_FILE = "faiss_index"
DATA_FILE = "data.json"

# Ensure Upload Directory Exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

# FAISS Index Setup
d = 384  # Embedding dimension
index = faiss.IndexFlatL2(d) if not os.path.exists(INDEX_FILE) else faiss.read_index(INDEX_FILE)

# Load Metadata
if os.path.exists(DATA_FILE):
    with open(DATA_FILE, "r") as f:
        metadata_store = json.load(f)
else:
    metadata_store = {}

def extract_text(file_path: str) -> str:
    """Extracts text from PDF or DOCX files."""
    _, ext = os.path.splitext(file_path)
    
    if ext.lower() == ".pdf":
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error reading PDF: {str(e)}")
    
    elif ext.lower() == ".docx":
        try:
            doc = docx.Document(file_path)
            return " ".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error reading DOCX: {str(e)}")
    
    raise HTTPException(status_code=400, detail="Unsupported file format")

@app.get("/query/")
async def query_documents(query: str = Query(..., min_length=3)):
    """Generates a meaningful answer based on document knowledge using Phi-2."""
    if not metadata_store:
        raise HTTPException(status_code=404, detail="No documents available for search")

    try:
        query_embedding = embedding_model.encode([query])[0].reshape(1, -1)
        distances, indices = index.search(query_embedding, 5)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

    best_match, best_file, best_content = None, None, None
    
    for dist, idx in zip(distances[0], indices[0]):
        if idx < len(metadata_store):
            file_id = list(metadata_store.keys())[idx]
            content = metadata_store[file_id]["content"]
            best_match, best_file, best_content = file_id, metadata_store[file_id]["filename"], content
            break  
    
    if not best_match:
        return JSONResponse(content={"message": "No relevant information found."})

    # Generate a meaningful answer using Phi-2
    prompt = f"Q: {query}\nBased on the following text, generate a human-like response:\n\n{best_content}\n\nA:"
    ai_response = qa_pipeline(prompt, max_length=100, num_return_sequences=1)[0]["generated_text"]
    
    return JSONResponse(content={"file": best_file, "answer": ai_response})
